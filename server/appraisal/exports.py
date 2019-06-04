from cornice.resource import resource
from pyramid.authorization import Allow, Everyone
from pyramid.response import Response
import tempfile
import subprocess
import os
import io
import json
from base64 import b64encode
import re
import requests
import base64
import bson
from PIL import Image
import hashlib
from pprint import pprint
import concurrent.futures
import filetype
from appraisal.components.document_processor import DocumentProcessor
from appraisal.models.comparable_sale import ComparableSale
from appraisal.models.comparable_lease import ComparableLease
from appraisal.models.appraisal import Appraisal
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.dml import MSO_THEME_COLOR
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import docx
from pyramid.security import Authenticated
from pyramid.authorization import Allow, Deny, Everyone
from appraisal.authorization import checkUserOwnsObject, getAccessTokenForRequest
from pyramid.httpexceptions import HTTPForbidden


class ExportAPI(object):
    def addBackground(self, cell, color):
        shading_elm_1 = parse_xml((r'<w:shd {} w:fill="{}" w:textColor="ffffff"/>').format(nsdecls('w'), color))
        cell._tc.get_or_add_tcPr().append(shading_elm_1)

    def addFontColor(self, cell, r, g, b):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = docx.shared.RGBColor(r, g, b)

    def formatAmount(self, number):
        number = "{:.2f}".format(number)
        regex = re.compile(r"\B(?=(\d{3})+(?!\d))")
        return regex.sub(", ", number)


    def renderTemplate(self, templateName, data):
        wordDocFolder = os.path.join(os.getcwd(), "appraisal", "word_documents")

        with tempfile.TemporaryDirectory() as tempDir:
            with tempfile.NamedTemporaryFile(suffix=".html", dir=tempDir) as temp:
                data["fileName"] = temp.name

                renderHtml = subprocess.run(args=["npm",  "run-script", templateName], cwd=wordDocFolder, input=bytes(json.dumps(data), 'utf8'), stdout=subprocess.PIPE, stderr=subprocess.PIPE)

                stdout = str(renderHtml.stdout, 'utf8')
                stderr = str(renderHtml.stderr, 'utf8')

                if renderHtml.returncode != 0 or stderr != "":
                    raise Exception(stderr)

                with open("data.html", "wt") as testFile:
                    testFile.write(open(temp.name, "rt").read())

                odtFileName = temp.name.replace(".html", ".odt")
                docxFileName = temp.name.replace(".html", ".docx")

                convertOdt = subprocess.run(args=["libreoffice", "--norestore", "--headless", "--convert-to", "odt", temp.name], cwd=tempDir, stdout=subprocess.PIPE, stderr=subprocess.PIPE)

                stdout = str(convertOdt.stdout, 'utf8')
                stderr = str(convertOdt.stderr, 'utf8')

                if convertOdt.returncode != 0:
                    raise Exception(stderr)


                convertDocx = subprocess.run(args=["libreoffice", "--norestore", "--headless", "--convert-to", "docx", odtFileName], cwd=tempDir, stdout=subprocess.PIPE, stderr=subprocess.PIPE)

                stdout = str(convertDocx.stdout, 'utf8')
                stderr = str(convertDocx.stderr, 'utf8')

                if convertDocx.returncode != 0:
                    raise Exception(stderr)

                file = open(docxFileName, 'rb')

                buffer = io.BytesIO()
                buffer.write(file.read())
                buffer.seek(0)

                file.close()

                os.unlink(odtFileName)
                os.unlink(docxFileName)

                return buffer

@resource(path='/appraisal/{appraisalId}/comparable_sales/excel', cors_enabled=True, cors_origins="*", permission="everything")
class ComparableSalesExcelFile(ExportAPI):

    def __init__(self, request, context=None):
        self.request = request

    def __acl__(self):
        return [
            (Allow, Authenticated, 'everything'),
            (Deny, Everyone, 'everything')
        ]

    def get(self):
        appraisalId = self.request.matchdict['appraisalId']

        appraisal = Appraisal.objects(id=appraisalId).first()

        auth = checkUserOwnsObject(self.request.authenticated_userid, self.request.effective_principals, appraisal)
        if not auth:
            raise HTTPForbidden("You do not have access to this appraisal.")

        query = {"id__in": appraisal.comparableSalesDCA + appraisal.comparableSalesCapRate}

        if "view_all" not in self.request.effective_principals:
            query["owner"] = self.request.authenticated_userid

        comparables = ComparableSale.objects(**query)

        wb = Workbook()

        ws1 = wb.active
        ws1.title = "Comparable Sales"

        ws1.append([
            "Name",
            "Address",
            "Sale Date",
            "Net Operating Income",
            "Sale Price",
            "Capitalization Rate",
            "Size (sf)",
            "TMI (psf)",
            "Vacancy Rate",
            "Description"
        ])

        for comp in comparables:
            ws1.append([
                comp.name,
                comp.address,
                comp.saleDate,
                comp.netOperatingIncome,
                comp.salePrice,
                comp.capitalizationRate,
                comp.sizeSquareFootage,
                comp.taxesMaintenanceInsurancePSF,
                comp.vacancyRate,
                comp.description,
            ])

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        response = Response()
        response.body_file = buffer
        response.content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        response.content_disposition = "attachment; filename=\"ComparableSales.xlsx\""
        return response




@resource(path='/appraisal/{appraisalId}/comparable_sales/word', cors_enabled=True, cors_origins="*", permission="everything")
class ComparableSalesWordFile(ExportAPI):

    def __init__(self, request, context=None):
        self.request = request

    def __acl__(self):
        return [
            (Allow, Authenticated, 'everything'),
            (Deny, Everyone, 'everything')
        ]

    def get(self):
        appraisalId = self.request.matchdict['appraisalId']

        appraisal = Appraisal.objects(id=appraisalId).first()

        auth = checkUserOwnsObject(self.request.authenticated_userid, self.request.effective_principals, appraisal)
        if not auth:
            raise HTTPForbidden("You do not have access to this appraisal.")

        query = {"id__in": appraisal.comparableSalesDCA + appraisal.comparableSalesCapRate}

        if "view_all" not in self.request.effective_principals:
            query["owner"] = self.request.authenticated_userid

        comparables = ComparableSale.objects(**query)

        data = {
            "appraisal": json.loads(appraisal.to_json()),
            "comparableSales": [json.loads(comp.to_json()) for comp in comparables],
            "accessToken": getAccessTokenForRequest(self.request)
        }

        buffer = self.renderTemplate("comparable_sales_summary_word", data)

        response = Response()
        response.body_file = buffer
        response.content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        response.content_disposition = "attachment; filename=\"ComparableSales.docx\""
        return response



@resource(path='/appraisal/{appraisalId}/comparable_sales/detailed_word', cors_enabled=True, cors_origins="*", permission="everything")
class ComparableSalesDetailedWordFile(ExportAPI):

    def __init__(self, request, context=None):
        self.request = request

    def __acl__(self):
        return [
            (Allow, Authenticated, 'everything'),
            (Deny, Everyone, 'everything')
        ]


    def addHeaderFormatting(self, cell):
        self.addFontColor(cell, 0, 0, 0)

    def get(self):
        appraisalId = self.request.matchdict['appraisalId']

        appraisal = Appraisal.objects(id=appraisalId).first()

        auth = checkUserOwnsObject(self.request.authenticated_userid, self.request.effective_principals, appraisal)
        if not auth:
            raise HTTPForbidden("You do not have access to this appraisal.")

        query = {"id__in": appraisal.comparableSalesDCA + appraisal.comparableSalesCapRate}

        if "view_all" not in self.request.effective_principals:
            query["owner"] = self.request.authenticated_userid

        comparables = ComparableSale.objects(**query)

        comparables = [comp for comp in comparables]

        accessToken = getAccessTokenForRequest(self.request)

        for comp in comparables:
            newUrls = []
            if len(comp.imageUrls) == 0:
                image = requests.get(f"https://maps.googleapis.com/maps/api/streetview?key=AIzaSyBRmZ2N4EhJjXmC29t3VeiLUQssNG-MY1I&size=640x480&source=outdoor&location={comp.address}").content
                data64 = u''.join(str(base64.encodebytes(image), 'utf8'))
                newUrls.append(u'data:%s;base64,%s' % ("image/jpeg", data64))
            else:
                for imageUrl in comp.imageUrls:
                    image = requests.get(imageUrl + f"?access_token=" + accessToken).content

                    data64 = u''.join(str(base64.encodebytes(image), 'utf8'))
                    newUrls.append(u'data:%s;base64,%s' % ("image/jpeg", data64))
                    break
            comp.imageUrls = newUrls

        data = {
            "appraisal": json.loads(appraisal.to_json()),
            "comparableSales": [json.loads(comp.to_json()) for comp in comparables]
        }

        buffer = self.renderTemplate("comparable_sales_detailed_word", data)

        response = Response()
        response.body_file = buffer
        response.content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        response.content_disposition = "attachment; filename=\"ComparableSales.docx\""
        return response




@resource(path='/appraisal/{appraisalId}/comparable_leases/excel', cors_enabled=True, cors_origins="*", permission="everything")
class ComparableLeasesExcelFile(ExportAPI):

    def __init__(self, request, context=None):
        self.request = request

    def __acl__(self):
        return [
            (Allow, Authenticated, 'everything'),
            (Deny, Everyone, 'everything')
        ]

    def get(self):
        appraisalId = self.request.matchdict['appraisalId']

        appraisal = Appraisal.objects(id=appraisalId).first()

        auth = checkUserOwnsObject(self.request.authenticated_userid, self.request.effective_principals, appraisal)
        if not auth:
            raise HTTPForbidden("You do not have access to this appraisal.")

        query = {"id__in": appraisal.comparableLeases}

        if "view_all" not in self.request.effective_principals:
            query["owner"] = self.request.authenticated_userid

        comparables = ComparableLease.objects(**query)

        wb = Workbook()

        ws1 = wb.active
        ws1.title = "Comparable Leases"

        ws1.append([
            "Name",
            "Address",
            "Lease Date",
            "Type",
            "Size (sf)",
            "Rent",
            "Description"
        ])

        for comp in comparables:
            ws1.append([
                comp.name,
                comp.address,
                comp.leaseDate,
                comp.propertyType,
                comp.sizeOfUnit,
                comp.yearlyRent,
                comp.description
            ])

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        response = Response()
        response.body_file = buffer
        response.content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        response.content_disposition = "attachment; filename=\"ComparableLeases.xlsx\""
        return response





@resource(path='/appraisal/{appraisalId}/comparable_leases/word', cors_enabled=True, cors_origins="*", permission="everything")
class ComparableLeasesWordFile(ExportAPI):

    def __init__(self, request, context=None):
        self.request = request


    def __acl__(self):
        return [
            (Allow, Authenticated, 'everything'),
            (Deny, Everyone, 'everything')
        ]

    def addHeaderFormatting(self, cell):
        self.addBackground(cell, "33339A")
        self.addFontColor(cell, 255, 255, 255)

    def get(self):
        appraisalId = self.request.matchdict['appraisalId']

        appraisal = Appraisal.objects(id=appraisalId).first()

        auth = checkUserOwnsObject(self.request.authenticated_userid, self.request.effective_principals, appraisal)
        if not auth:
            raise HTTPForbidden("You do not have access to this appraisal.")

        query = {"id__in": appraisal.comparableLeases}

        if "view_all" not in self.request.effective_principals:
            query["owner"] = self.request.authenticated_userid

        comparables = ComparableLease.objects(**query)

        document = Document()

        document.add_heading('Comparable Leases', 0)

        table = document.add_table(rows=1, cols=5)
        table.style = 'TableGrid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Date'
        hdr_cells[1].text = 'Address'
        hdr_cells[2].text = 'Area'
        hdr_cells[3].text = 'Rent'
        hdr_cells[4].text = 'Remarks'
        for cell in hdr_cells:
            self.addHeaderFormatting(cell)

        document.add_page_break()

        document.save('demo.docx')

        for comp in comparables:
            row_cells = table.add_row().cells
            if comp.leaseDate is not None:
                row_cells[0].text = comp.leaseDate.strftime("%m/%y")
            if comp.address is not None:
                row_cells[1].text = comp.address
            if comp.sizeOfUnit is not None:
                row_cells[2].text = self.formatAmount(comp.sizeOfUnit) + " sf"
            if comp.rentEscalations and len(comp.rentEscalations) > 0 and comp.rentEscalations[0].yearlyRent is not None and comp.sizeOfUnit is not None:
                row_cells[3].text = "$" + self.formatAmount(comp.rentEscalations[0].yearlyRent / comp.sizeOfUnit)
            if comp.description is not None:
                row_cells[4].text = comp.description

            for cell in row_cells:
                self.addBackground(cell, "FFFDF3")

        table.columns[0].width = Inches(0.75)
        table.columns[1].width = Inches(1.5)
        table.columns[2].width = Inches(1)
        table.columns[3].width = Inches(0.75)
        table.columns[4].width = Inches(2.0)

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)

        response = Response()
        response.body_file = buffer
        response.content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        response.content_disposition = "attachment; filename=\"ComparableLeases.docx\""
        return response



@resource(path='/appraisal/{appraisalId}/rent_roll/word', cors_enabled=True, cors_origins="*", permission="everything")
class RentRollWordFile(ExportAPI):

    def __init__(self, request, context=None):
        self.request = request

    def __acl__(self):
        return [
            (Allow, Authenticated, 'everything'),
            (Deny, Everyone, 'everything')
        ]


    def addHeaderFormatting(self, cell):
        self.addBackground(cell, "33339A")
        self.addFontColor(cell, 255, 255, 255)

    def get(self):
        appraisalId = self.request.matchdict['appraisalId']

        appraisal = Appraisal.objects(id=appraisalId).first()

        auth = checkUserOwnsObject(self.request.authenticated_userid, self.request.effective_principals, appraisal)
        if not auth:
            raise HTTPForbidden("You do not have access to this appraisal.")

        data = {
            "appraisal": json.loads(appraisal.to_json())
        }

        for unitIndex, unit in enumerate(data['appraisal']['units']):
            unit['currentTenancy'] = json.loads(appraisal.units[unitIndex].currentTenancy.to_json())

        buffer = self.renderTemplate("rent_roll_summary_word", data)

        response = Response()
        response.body_file = buffer
        response.content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        response.content_disposition = "attachment; filename=\"RentRoll.docx\""
        return response



@resource(path='/appraisal/{appraisalId}/rent_roll/excel', cors_enabled=True, cors_origins="*", permission="everything")
class RentRollExcelFile(ExportAPI):

    def __init__(self, request, context=None):
        self.request = request

    def __acl__(self):
        return [
            (Allow, Authenticated, 'everything'),
            (Deny, Everyone, 'everything')
        ]

    def get(self):
        appraisalId = self.request.matchdict['appraisalId']

        appraisal = Appraisal.objects(id=appraisalId).first()

        auth = checkUserOwnsObject(self.request.authenticated_userid, self.request.effective_principals, appraisal)
        if not auth:
            raise HTTPForbidden("You do not have access to this appraisal.")

        wb = Workbook()

        ws1 = wb.active
        ws1.title = "Rent Roll"

        ws1.append([
            "Unit Number",
            "Size (sqft)",
            "Tenant Name",
            "Yearly Rent ($)",
            "Remarks"
        ])

        for unit in appraisal.units:
            name = ""
            if unit.currentTenancy is not None:
                name = unit.currentTenancy.name

            yearlyRent = ""
            if unit.currentTenancy is not None:
                yearlyRent = unit.currentTenancy.yearlyRent

            ws1.append([
                unit.unitNumber,
                unit.squareFootage,
                name,
                yearlyRent,
                unit.remarks
            ])

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        response = Response()
        response.body_file = buffer
        response.content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        response.content_disposition = "attachment; filename=\"RentRoll.xlsx\""
        return response



@resource(path='/appraisal/{appraisalId}/expenses/word', cors_enabled=True, cors_origins="*", permission="everything")
class ExpensesWordFile(ExportAPI):

    def __init__(self, request, context=None):
        self.request = request

    def __acl__(self):
        return [
            (Allow, Authenticated, 'everything'),
            (Deny, Everyone, 'everything')
        ]

    def get(self):
        appraisalId = self.request.matchdict['appraisalId']

        appraisal = Appraisal.objects(id=appraisalId).first()

        auth = checkUserOwnsObject(self.request.authenticated_userid, self.request.effective_principals, appraisal)
        if not auth:
            raise HTTPForbidden("You do not have access to this appraisal.")

        data = {
            "appraisal": json.loads(appraisal.to_json())
        }

        buffer = self.renderTemplate("expenses_summary_word", data)

        response = Response()
        response.body_file = buffer
        response.content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        response.content_disposition = "attachment; filename=\"Expenses.docx\""
        return response


@resource(path='/appraisal/{appraisalId}/expenses/excel', cors_enabled=True, cors_origins="*", permission="everything")
class ExpensesExcelFile(ExportAPI):

    def __init__(self, request, context=None):
        self.request = request

    def __acl__(self):
        return [
            (Allow, Authenticated, 'everything'),
            (Deny, Everyone, 'everything')
        ]
    def get(self):
        appraisalId = self.request.matchdict['appraisalId']

        appraisal = Appraisal.objects(id=appraisalId).first()

        auth = checkUserOwnsObject(self.request.authenticated_userid, self.request.effective_principals, appraisal)
        if not auth:
            raise HTTPForbidden("You do not have access to this appraisal.")

        wb = Workbook()

        ws1 = wb.active
        ws1.title = "Expenses"

        headers = [""]

        for year in appraisal.incomeStatement.years:
            headers.append(str(year))

        headers[0] = "Operating Expenses"
        ws1.append(headers)

        operatingExpenses = [expense for expense in appraisal.incomeStatement.expenses if expense.incomeStatementItemType == 'operating_expense']
        managementExpenses = [expense for expense in appraisal.incomeStatement.expenses if expense.incomeStatementItemType == 'management_expense']
        taxExpenses = [expense for expense in appraisal.incomeStatement.expenses if expense.incomeStatementItemType == 'taxes']

        for expense in operatingExpenses:
            row = [expense.name]

            for year in appraisal.incomeStatement.years:
                row.append(expense.yearlyAmounts.get(str(year)))

            ws1.append(row)

        ws1.append([""] * len(headers))
        ws1.append([""] * len(headers))
        headers[0] = "Management Expenses"
        ws1.append(headers)

        for expense in managementExpenses:
            row = [expense.name]

            for year in appraisal.incomeStatement.years:
                row.append(expense.yearlyAmounts.get(str(year)))

            ws1.append(row)

        ws1.append([""] * len(headers))
        ws1.append([""] * len(headers))
        headers[0] = "Tax Expenses"
        ws1.append(headers)

        for expense in taxExpenses:
            row = [expense.name]

            for year in appraisal.incomeStatement.years:
                row.append(expense.yearlyAmounts.get(str(year)))

            ws1.append(row)

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        response = Response()
        response.body_file = buffer
        response.content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        response.content_disposition = "attachment; filename=\"Expenses.xlsx\""
        return response


@resource(path='/appraisal/{appraisalId}/stabilized_statement/word', cors_enabled=True, cors_origins="*", permission="everything")
class StabilizedStatementWordFile(ExportAPI):

    def __init__(self, request, context=None):
        self.request = request

    def __acl__(self):
        return [
            (Allow, Authenticated, 'everything'),
            (Deny, Everyone, 'everything')
        ]

    def get(self):
        appraisalId = self.request.matchdict['appraisalId']

        appraisal = Appraisal.objects(id=appraisalId).first()

        auth = checkUserOwnsObject(self.request.authenticated_userid, self.request.effective_principals, appraisal)
        if not auth:
            raise HTTPForbidden("You do not have access to this appraisal.")

        data = {
            "appraisal": json.loads(appraisal.to_json())
        }

        buffer = self.renderTemplate("stabilized_statement_word", data)

        response = Response()
        response.body_file = buffer
        response.content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        response.content_disposition = "attachment; filename=\"StabilizedStatement.docx\""
        return response


@resource(path='/appraisal/{appraisalId}/stabilized_statement/excel', cors_enabled=True, cors_origins="*", permission="everything")
class StabilizedStatementExcelFile(ExportAPI):

    def __init__(self, request, context=None):
        self.request = request

    def __acl__(self):
        return [
            (Allow, Authenticated, 'everything'),
            (Deny, Everyone, 'everything')
        ]

    def get(self):
        appraisalId = self.request.matchdict['appraisalId']

        appraisal = Appraisal.objects(id=appraisalId).first()

        auth = checkUserOwnsObject(self.request.authenticated_userid, self.request.effective_principals, appraisal)
        if not auth:
            raise HTTPForbidden("You do not have access to this appraisal.")

        wb = Workbook()

        ws1 = wb.active
        ws1.title = "Stabilized Statement"

        headers = ["Stabilized Statement"]

        ws1.append(headers)

        ws1.append(["Revenue", "", ""])

        ws1.append(["Stabilized Rental Income", appraisal.stabilizedStatement.rentalIncome, ""])

        if appraisal.stabilizedStatement.additionalIncome:
            ws1.append(["Additional Income", appraisal.stabilizedStatement.additionalIncome, ""])

        ws1.append(["Recoverable Income", appraisal.stabilizedStatement.recoverableIncome, ""])

        ws1.append(["Potential Gross Income", appraisal.stabilizedStatement.potentialGrossIncome, ""])

        ws1.append([f"Less Vacancy @ {appraisal.stabilizedStatementInputs.vacancyRate}", appraisal.stabilizedStatement.vacancyDeduction, ""])

        ws1.append([f"Effective Gross Income", "", appraisal.stabilizedStatement.effectiveGrossIncome])

        ws1.append([""])

        ws1.append(["Expenses"])

        if appraisal.stabilizedStatement.operatingExpenses:
            ws1.append(["Operating Expenses", appraisal.stabilizedStatement.operatingExpenses, ""])

        if appraisal.stabilizedStatement.taxes:
            ws1.append(["Taxes", appraisal.stabilizedStatement.taxes, ""])

        if appraisal.stabilizedStatement.managementExpenses:
            label = "Management Expenses"

            if appraisal.stabilizedStatementInputs.managementExpenseMode == 'combined_structural_rule':
                label = "Structural & Mgmt"

            ws1.append([label, appraisal.stabilizedStatement.managementExpenses, ""])

        if appraisal.stabilizedStatement.tmiTotal:
            ws1.append([f"TMI {appraisal.sizeOfBuilding} sqft @ ${appraisal.stabilizedStatementInputs.tmiRatePSF}", appraisal.stabilizedStatement.tmiTotal, ""])

        if appraisal.stabilizedStatement.structuralAllowance:
            ws1.append([f"Structural Allowance @ {appraisal.stabilizedStatementInputs.structuralAllowancePercent}%", appraisal.stabilizedStatement.structuralAllowance, ""])


        ws1.append([f"Total Expenses", "", appraisal.stabilizedStatement.totalExpenses])

        ws1.append([f"Net Operating Income", "", appraisal.stabilizedStatement.netOperatingIncome])

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        response = Response()
        response.body_file = buffer
        response.content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        response.content_disposition = "attachment; filename=\"StabilizedStatement.xlsx\""
        return response



@resource(path='/appraisal/{appraisalId}/capitalization_valuation/word', cors_enabled=True, cors_origins="*", permission="everything")
class CapitalizationValuationWordFile(ExportAPI):

    def __init__(self, request, context=None):
        self.request = request

    def __acl__(self):
        return [
            (Allow, Authenticated, 'everything'),
            (Deny, Everyone, 'everything')
        ]

    def get(self):
        appraisalId = self.request.matchdict['appraisalId']

        appraisal = Appraisal.objects(id=appraisalId).first()

        auth = checkUserOwnsObject(self.request.authenticated_userid, self.request.effective_principals, appraisal)
        if not auth:
            raise HTTPForbidden("You do not have access to this appraisal.")

        data = {
            "appraisal": json.loads(appraisal.to_json())
        }

        buffer = self.renderTemplate("capitalization_valuation_word", data)

        response = Response()
        response.body_file = buffer
        response.content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        response.content_disposition = "attachment; filename=\"CapitalizationValuation.docx\""
        return response


@resource(path='/appraisal/{appraisalId}/capitalization_valuation/excel', cors_enabled=True, cors_origins="*", permission="everything")
class CapitalizationValuationExcelFile(ExportAPI):

    def __init__(self, request, context=None):
        self.request = request

    def __acl__(self):
        return [
            (Allow, Authenticated, 'everything'),
            (Deny, Everyone, 'everything')
        ]

    def get(self):
        appraisalId = self.request.matchdict['appraisalId']

        appraisal = Appraisal.objects(id=appraisalId).first()

        auth = checkUserOwnsObject(self.request.authenticated_userid, self.request.effective_principals, appraisal)
        if not auth:
            raise HTTPForbidden("You do not have access to this appraisal.")

        wb = Workbook()

        ws1 = wb.active
        ws1.title = "Capitalization Valuation"

        headers = ["Capitalization Valuation"]

        ws1.append(headers)

        ws1.append(["Net Operating Income", appraisal.stabilizedStatement.netOperatingIncome, ""])

        ws1.append([f"Capitalized @ {appraisal.stabilizedStatementInputs.capitalizationRate}%", appraisal.stabilizedStatement.capitalization, ""])

        if appraisal.stabilizedStatement.marketRentDifferential and appraisal.stabilizedStatementInputs.applyMarketRentDifferential:
            ws1.append(["Market Rent Differential", appraisal.stabilizedStatement.marketRentDifferential, ""])

        if appraisal.stabilizedStatement.vacantUnitRentLoss and appraisal.stabilizedStatementInputs.applyVacantUnitRentLoss:
            ws1.append(["Vacant Unit Rent Loss", appraisal.stabilizedStatement.vacantUnitRentLoss, ""])

        if appraisal.stabilizedStatement.freeRentRentLoss and appraisal.stabilizedStatementInputs.applyFreeRentLoss:
            ws1.append(["Free Rent Loss", appraisal.stabilizedStatement.freeRentRentLoss, ""])

        if appraisal.stabilizedStatement.vacantUnitLeasupCosts and appraisal.stabilizedStatementInputs.applyVacantUnitLeasingCosts:
            ws1.append(["Vacant Unit Leasup Costs", appraisal.stabilizedStatement.vacantUnitLeasupCosts, ""])

        if appraisal.stabilizedStatement.amortizedCapitalInvestment and appraisal.stabilizedStatementInputs.applyAmortization:
            ws1.append(["Amortized Capital Investment", appraisal.stabilizedStatement.amortizedCapitalInvestment, ""])

        for modifier in appraisal.stabilizedStatementInputs.modifiers:
            ws1.append([modifier.name, modifier.amount, ""])

        ws1.append([f"Valuation", "", appraisal.stabilizedStatement.valuation])

        ws1.append([f"Rounded", "", appraisal.stabilizedStatement.valuationRounded])

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        response = Response()
        response.body_file = buffer
        response.content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        response.content_disposition = "attachment; filename=\"CapitalizationValuation.xlsx\""
        return response





@resource(path='/appraisal/{appraisalId}/direct_comparison_valuation/word', cors_enabled=True, cors_origins="*", permission="everything")
class DirectComparisonValuationWordFile(ExportAPI):

    def __init__(self, request, context=None):
        self.request = request

    def __acl__(self):
        return [
            (Allow, Authenticated, 'everything'),
            (Deny, Everyone, 'everything')
        ]

    def get(self):
        appraisalId = self.request.matchdict['appraisalId']

        appraisal = Appraisal.objects(id=appraisalId).first()

        auth = checkUserOwnsObject(self.request.authenticated_userid, self.request.effective_principals, appraisal)
        if not auth:
            raise HTTPForbidden("You do not have access to this appraisal.")

        data = {
            "appraisal": json.loads(appraisal.to_json())
        }

        buffer = self.renderTemplate("direct_comparison_valuation_word", data)

        response = Response()
        response.body_file = buffer
        response.content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        response.content_disposition = "attachment; filename=\"DirectComparisonValuation.docx\""
        return response



@resource(path='/appraisal/{appraisalId}/direct_comparison_valuation/excel', cors_enabled=True, cors_origins="*", permission="everything")
class DirectComparisonValuationExcelFile(ExportAPI):

    def __init__(self, request, context=None):
        self.request = request

    def __acl__(self):
        return [
            (Allow, Authenticated, 'everything'),
            (Deny, Everyone, 'everything')
        ]

    def get(self):
        appraisalId = self.request.matchdict['appraisalId']

        appraisal = Appraisal.objects(id=appraisalId).first()

        auth = checkUserOwnsObject(self.request.authenticated_userid, self.request.effective_principals, appraisal)
        if not auth:
            raise HTTPForbidden("You do not have access to this appraisal.")

        wb = Workbook()

        ws1 = wb.active
        ws1.title = "Direct Comparison Valuation"

        headers = ["Direct Comparison Valuation"]

        ws1.append(headers)

        if appraisal.directComparisonInputs.directComparisonMetric == 'psf':
            ws1.append([f"{appraisal.sizeOfBuilding} sqft @ {appraisal.directComparisonInputs.pricePerSquareFoot}%", appraisal.directComparisonValuation.comparativeValue, ""])
        elif appraisal.directComparisonInputs.directComparisonMetric == 'noi_multiple':
            ws1.append([f"{appraisal.sizeOfBuilding} sqft @ ${appraisal.directComparisonInputs.noiPSFPricePerSquareFoot}", appraisal.directComparisonValuation.comparativeValue, ""])
        elif appraisal.directComparisonInputs.directComparisonMetric == 'psf_land':
            ws1.append([f"{appraisal.sizeOfLand * 43560} sqft @ ${appraisal.directComparisonInputs.pricePerSquareFootLand}", appraisal.directComparisonValuation.comparativeValue, ""])
        elif appraisal.directComparisonInputs.directComparisonMetric == 'per_acre_land':
            ws1.append([f"{appraisal.sizeOfLand} acres @ ${appraisal.directComparisonInputs.pricePerAcreLand}", appraisal.directComparisonValuation.comparativeValue, ""])
        elif appraisal.directComparisonInputs.directComparisonMetric == 'psf_buildable_area':
            ws1.append([f"{appraisal.buildableArea} sqft @ ${appraisal.directComparisonInputs.pricePerSquareFootBuildableArea}", appraisal.directComparisonValuation.comparativeValue, ""])
        elif appraisal.directComparisonInputs.directComparisonMetric == 'psf_buildable_area':
            ws1.append([f"{appraisal.buildableUnits} units @ ${appraisal.directComparisonInputs.pricePerBuildableUnit}", appraisal.directComparisonValuation.comparativeValue, ""])

        if appraisal.directComparisonValuation.marketRentDifferential and appraisal.directComparisonInputs.applyMarketRentDifferential:
            ws1.append(["Market Rent Differential", appraisal.directComparisonValuation.marketRentDifferential, ""])

        if appraisal.directComparisonValuation.vacantUnitRentLoss and appraisal.directComparisonInputs.applyVacantUnitRentLoss:
            ws1.append(["Vacant Unit Rent Loss", appraisal.directComparisonValuation.vacantUnitRentLoss, ""])

        if appraisal.directComparisonValuation.freeRentRentLoss and appraisal.directComparisonInputs.applyFreeRentLoss:
            ws1.append(["Free Rent Loss", appraisal.directComparisonValuation.freeRentRentLoss, ""])

        if appraisal.directComparisonValuation.vacantUnitLeasupCosts and appraisal.directComparisonInputs.applyVacantUnitLeasingCosts:
            ws1.append(["Vacant Unit Leasup Costs", appraisal.directComparisonValuation.vacantUnitLeasupCosts, ""])

        if appraisal.directComparisonValuation.amortizedCapitalInvestment and appraisal.directComparisonInputs.applyAmortization:
            ws1.append(["Amortized Capital Investment", appraisal.directComparisonValuation.amortizedCapitalInvestment, ""])

        for modifier in appraisal.directComparisonInputs.modifiers:
            ws1.append([modifier.name, modifier.amount, ""])

        ws1.append([f"Valuation", "", appraisal.directComparisonValuation.valuation])

        ws1.append([f"Rounded", "", appraisal.directComparisonValuation.valuationRounded])

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        response = Response()
        response.body_file = buffer
        response.content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        response.content_disposition = "attachment; filename=\"DirectComparisonValuation.xlsx\""
        return response




@resource(path='/appraisal/{appraisalId}/additional_incomes/word', cors_enabled=True, cors_origins="*", permission="everything")
class AdditionalIncomeWordFile(ExportAPI):

    def __init__(self, request, context=None):
        self.request = request

    def __acl__(self):
        return [
            (Allow, Authenticated, 'everything'),
            (Deny, Everyone, 'everything')
        ]

    def get(self):
        appraisalId = self.request.matchdict['appraisalId']

        appraisal = Appraisal.objects(id=appraisalId).first()

        auth = checkUserOwnsObject(self.request.authenticated_userid, self.request.effective_principals, appraisal)
        if not auth:
            raise HTTPForbidden("You do not have access to this appraisal.")

        data = {
            "appraisal": json.loads(appraisal.to_json())
        }

        buffer = self.renderTemplate("additional_income_word", data)

        response = Response()
        response.body_file = buffer
        response.content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        response.content_disposition = "attachment; filename=\"AdditionalIncome.docx\""
        return response



@resource(path='/appraisal/{appraisalId}/additional_incomes/excel', cors_enabled=True, cors_origins="*", permission="everything")
class AdditionalIncomeExcelFile(ExportAPI):

    def __init__(self, request, context=None):
        self.request = request

    def __acl__(self):
        return [
            (Allow, Authenticated, 'everything'),
            (Deny, Everyone, 'everything')
        ]
    def get(self):
        appraisalId = self.request.matchdict['appraisalId']

        appraisal = Appraisal.objects(id=appraisalId).first()

        auth = checkUserOwnsObject(self.request.authenticated_userid, self.request.effective_principals, appraisal)
        if not auth:
            raise HTTPForbidden("You do not have access to this appraisal.")

        wb = Workbook()

        ws1 = wb.active
        ws1.title = "Additional Income"

        headers = [""]

        for year in appraisal.incomeStatement.years:
            headers.append(str(year))

        headers[0] = "Additional Income"
        ws1.append(headers)

        operatingIncomes = [income for income in appraisal.incomeStatement.incomes if income.incomeStatementItemType == 'additional_income']

        for income in operatingIncomes:
            row = [income.name]

            for year in appraisal.incomeStatement.years:
                row.append(income.yearlyAmounts.get(str(year)))

            ws1.append(row)

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        response = Response()
        response.body_file = buffer
        response.content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        response.content_disposition = "attachment; filename=\"AdditionalIncome.xlsx\""
        return response



@resource(path='/appraisal/{appraisalId}/amortization_schedule/word', cors_enabled=True, cors_origins="*", permission="everything")
class AmortizationScheduleWordFile(ExportAPI):

    def __init__(self, request, context=None):
        self.request = request

    def __acl__(self):
        return [
            (Allow, Authenticated, 'everything'),
            (Deny, Everyone, 'everything')
        ]

    def get(self):
        appraisalId = self.request.matchdict['appraisalId']

        appraisal = Appraisal.objects(id=appraisalId).first()

        auth = checkUserOwnsObject(self.request.authenticated_userid, self.request.effective_principals, appraisal)
        if not auth:
            raise HTTPForbidden("You do not have access to this appraisal.")

        data = {
            "appraisal": json.loads(appraisal.to_json())
        }

        buffer = self.renderTemplate("amortization_schedule_word", data)

        response = Response()
        response.body_file = buffer
        response.content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        response.content_disposition = "attachment; filename=\"AmortizationSchedule.docx\""
        return response



@resource(path='/appraisal/{appraisalId}/amortization_schedule/excel', cors_enabled=True, cors_origins="*", permission="everything")
class AmortizationScheduleExcelFile(ExportAPI):

    def __init__(self, request, context=None):
        self.request = request

    def __acl__(self):
        return [
            (Allow, Authenticated, 'everything'),
            (Deny, Everyone, 'everything')
        ]
    def get(self):
        appraisalId = self.request.matchdict['appraisalId']

        appraisal = Appraisal.objects(id=appraisalId).first()

        auth = checkUserOwnsObject(self.request.authenticated_userid, self.request.effective_principals, appraisal)
        if not auth:
            raise HTTPForbidden("You do not have access to this appraisal.")

        wb = Workbook()

        ws1 = wb.active
        ws1.title = "Amortization Schedule"

        headers = ["Name", "Amount", "Interest", "Discount Rate", "Start Date", "Period (months)"]

        ws1.append(headers)

        for item in appraisal.amortizationSchedule.items:
            row = [
                item.name,
                item.amount,
                item.interest,
                item.discountRate,
                str(item.startDate),
                item.periodMonths
            ]

            ws1.append(row)

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        response = Response()
        response.body_file = buffer
        response.content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        response.content_disposition = "attachment; filename=\"AmortizationSchedule.xlsx\""
        return response

