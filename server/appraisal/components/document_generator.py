import pkg_resources
import csv
import random
import math
from pprint import pprint
import docx
import io
from .document_parser import DocumentParser

class DocumentGenerator:
    def __init__(self):
        self.templates = [file for file in pkg_resources.resource_listdir("appraisal", "templates") if file.endswith('.docx')]

        self.templateFieldFiles = {
            file.replace(".csv", ""): file
            for file in pkg_resources.resource_listdir("appraisal", "templates")
            if file.endswith(".csv")
        }

        self.parser = DocumentParser()


    def loadTemplateFields(self, templateType):
        fieldsFile = self.templateFieldFiles[templateType]
        csvFileStream = pkg_resources.resource_stream('appraisal', f'templates/{fieldsFile}')

        rows = csv.DictReader(io.TextIOWrapper(csvFileStream))

        fields = {}

        for row in rows:
            for key in row:
                if row[key] is not None and row[key] != "":
                    if key not in fields:
                        fields[key] = []
                    fields[key].append(row[key])
        return fields


    def generateDocument(self, template, templateType):
        fields = self.loadTemplateFields(templateType)

        # First we go through the document, and replace any fields with IDS
        document = docx.Document(pkg_resources.resource_stream('appraisal', f'templates/{template}'))

        wordFields = {}
        currentId = 0

        paragraphs = [p for p in document.paragraphs]
        tables = [t for t in document.tables]
        for section in document.sections:
            paragraphs.extend(section.header.paragraphs)
            paragraphs.extend(section.footer.paragraphs)

            tables.extend(section.header.tables)
            tables.extend(section.footer.tables)

        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)

        for paragraph in paragraphs:
            for run in paragraph.runs:
                for field in fields:
                    keyname = f"<<{field}>>"

                    if keyname in run.text:
                        replacement = f"<<{currentId}>>"
                        wordFields[currentId] = field
                        currentId += 1

                        run.text = run.text.replace(keyname, replacement)

        bufferFile = io.BytesIO()
        document.save(bufferFile)

        # Now we convert the template into words with coordinates. This allows us to identify
        # the location of the template tokens within the original document, in terms of X/Y coordinates.
        images, templateWords = self.parser.processDocx(bufferFile.getbuffer())

        tokenWords = {}
        for word in templateWords:
            # print(word['word'])
            startString = "&lt;&lt;"
            endString = "&gt;&gt;"

            if startString in word['word'] and endString in word['word']:
                if word['word'] in tokenWords:
                    raise ValueError("Template has multiple injection points with same label.")

                startIndex = word['word'].find(startString)
                endIndex = word['word'].find(endString)

                replacementId = int(word['word'][startIndex+len(startString):endIndex])
                word['replacementId'] = replacementId

                tokenWords[replacementId] = word

        # Now go through the document a second time and replace those fields with a random replacement
        for paragraph in paragraphs:
            for run in paragraph.runs:
                while "<<" in run.text and ">>" in run.text:
                    replacementId = int(run.text[run.text.index('<<')+2:run.text.index(">>")])

                    keyname = f"<<{replacementId}>>"

                    if keyname in run.text:
                        replacement = random.choice(fields[wordFields[replacementId]])
                        run.text = run.text.replace(keyname, replacement)
                        tokenWords[replacementId]['replacement'] = replacement


        bufferFile = io.BytesIO()
        document.save(bufferFile)

        data = bufferFile.getbuffer()

        images, words = self.parser.processDocx(data)

        for replacementId in tokenWords:
            tokenWord = tokenWords[replacementId]
            replacementWords = tokenWord['replacement'].split()

            matchingWords = {
                self.removeSymbols(word): [] for word in replacementWords
            }

            for word in words:
                if self.removeSymbols(word['word']) in matchingWords:
                    matchingWords[self.removeSymbols(word['word'])].append(word)

            for replacementWord in matchingWords:
                replacementMatches = matchingWords[replacementWord]

                bestWord = None
                bestDistance = None
                for word in replacementMatches:
                    distX = abs(word['left'] - tokenWord['left'])
                    distY = abs((word['top'] + word['page']) - (tokenWord['top'] + tokenWord['page']))

                    dist = math.sqrt(distX * distX + distY * distY)

                    if bestDistance is None or dist < bestDistance:
                        bestWord = word

                if bestWord is not None:
                    bestWord['classification'] = wordFields[replacementId]
        # pprint(words)
        return {"words": words}

    def removeSymbols(self, text):
        symbols = ".,<>!@#$%^&*(){}[]|\;:'\"/?"
        for symbol in symbols:
            text = text.replace(symbol, "")
        return text