from __future__ import annotations

import csv
import hashlib
import json
import logging
import shutil
from io import BytesIO
from contextlib import asynccontextmanager
from datetime import UTC, datetime
from pathlib import Path
from uuid import uuid4

import fitz
from fastapi import BackgroundTasks, FastAPI, File as Upload, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from pymongo import MongoClient, ReturnDocument
from docx import Document
from openpyxl import Workbook

from .extraction import normalize_extraction, provider_for
from .calculations import refresh_valuations, unit_is_vacant, unit_tenant_name, unit_yearly_rent
from .schemas import ExtractionPatch, ExtractionResult
from .settings import settings

log = logging.getLogger("swiftly")
PUBLIC_WORD_FIELDS = {
    "word", "page", "lineNumber", "documentLineNumber", "column", "documentColumn",
    "index", "left", "right", "top", "bottom",
}

DEMO_DEFAULTS = {
    "appraisalType": "detailed", "units": [], "imageUrls": [], "captions": [], "propertyTags": [],
    "marketRents": [], "recoveryStructures": [], "leasingCosts": [], "comparableSalesCapRate": [],
    "comparableSalesDCA": [], "comparableLeases": [], "additionalIncomes": [], "expenses": [],
    "incomeStatement": {"years": [], "items": [], "incomes": [], "expenses": [], "yearlySourceTypes": {}},
    "expenseStatement": {"years": [], "items": [], "incomes": [], "expenses": [], "yearlySourceTypes": {}},
    "amortizationSchedule": {"items": []},
    "stabilizedStatementInputs": {"vacancyRate": 5, "capitalizationRate": 5, "structuralAllowancePercent": 0, "modifiers": []},
    "directComparisonInputs": {"directComparisonMetric": "psf", "modifiers": []},
    "discountedCashFlowInputs": {"projectionYears": 10, "discountRate": 0, "inflation": 0},
}

DEMO_SEED_VERSION = "2026.13"
DEMO_APPRAISAL = {
    "_id": "demo-appraisal", "owner": "local-demo", "name": "Harbour Centre Demo",
    "address": "100 Harbour Street, Toronto", "appraisalType": "detailed",
    "location": {"type": "Point", "coordinates": [-79.3777, 43.6426]},
    "propertyType": "office", "sizeOfBuilding": 15_000, "sizeOfLand": 1.4,
    "zoning": "CR 3.0", "effectiveDate": "2026-01-01",
    "units": [
        {"tenantName": "Northstar Foods", "suite": "101", "unitNumber": "101", "squareFootage": 8_000,
         "yearlyRent": 192_000, "tenancies": [{"name": "Northstar Foods", "yearlyRent": 192_000,
                                                  "startDate": "2024-01-01", "endDate": "2029-12-31"}]},
        {"tenantName": "Cedar Advisory", "suite": "205", "unitNumber": "205", "squareFootage": 5_000,
         "yearlyRent": 145_000, "tenancies": [{"name": "Cedar Advisory", "yearlyRent": 145_000,
                                                 "startDate": "2025-07-01", "endDate": "2031-06-30"}]},
        {"tenantName": "Vacant", "suite": "310", "unitNumber": "310", "squareFootage": 2_000,
         "marketRent": "Office market rent", "isVacantForStabilizedStatement": True, "shouldTreatAsVacant": True,
         "shouldUseMarketRent": True, "tenancies": [{"name": "Vacant", "yearlyRent": 0}]},
    ],
    "incomeStatement": {"years": [2026], "items": [
        {"name": "Parking", "incomeStatementItemType": "additional_income",
         "yearlyAmounts": {"2026": 42_000}, "yearlySourceTypes": {}, "extractionReferences": {}},
    ], "incomes": [], "expenses": [], "yearlySourceTypes": {}},
    "expenseStatement": {"years": [2026], "items": [
        {"name": "Repairs and maintenance", "incomeStatementItemType": "operating_expense",
         "yearlyAmounts": {"2026": 68_000}, "yearlySourceTypes": {}, "extractionReferences": {}},
        {"name": "Utilities", "incomeStatementItemType": "operating_expense",
         "yearlyAmounts": {"2026": 37_000}, "yearlySourceTypes": {}, "extractionReferences": {}},
        {"name": "Property taxes", "incomeStatementItemType": "taxes",
         "yearlyAmounts": {"2026": 82_000}, "yearlySourceTypes": {}, "extractionReferences": {}},
        {"name": "Management fee", "incomeStatementItemType": "management_expense",
         "yearlyAmounts": {"2026": 21_000}, "yearlySourceTypes": {}, "extractionReferences": {}},
    ], "incomes": [], "expenses": [], "yearlySourceTypes": {}},
    "additionalIncomes": [{"name": "Parking", "amount": 42_000}],
    "expenses": [{"name": "Repairs and maintenance", "amount": 68_000}, {"name": "Utilities", "amount": 37_000}],
    "marketRents": [{"name": "Office market rent", "amountPSF": 24, "amount": 24}],
    "amortizationSchedule": {"items": [{"name": "Tenant improvements", "amount": 18_000,
                                           "interest": 3, "discountRate": 7,
                                           "startDate": "2026-01-01", "periodMonths": 60}]},
    "stabilizedStatementInputs": {"vacancyRate": 5, "capitalizationRate": 5.25,
                                    "structuralAllowancePercent": 2, "expensesMode": "income_statement",
                                    "managementExpenseMode": "income_statement", "modifiers": []},
    "directComparisonInputs": {"directComparisonMetric": "psf", "pricePerSquareFoot": 260, "modifiers": []},
    "discountedCashFlowInputs": {"projectionYears": 10, "discountRate": 7, "inflation": 2},
    "dataTypeReferences": {
        "INCOME_STATEMENT": [{"appraisalId": "demo-appraisal", "fileId": "demo-financial-statement",
                              "pageNumbers": [1], "wordIndexes": []}],
        "EXPENSE_STATEMENT": [{"appraisalId": "demo-appraisal", "fileId": "demo-financial-statement",
                               "pageNumbers": [1], "wordIndexes": []}],
        "RENT_ROLL": [{"appraisalId": "demo-appraisal", "fileId": "demo-scanned-rent-roll",
                       "pageNumbers": [1], "wordIndexes": []}],
    },
    "comparableSalesCapRate": ["demo-sale"], "comparableSalesDCA": ["demo-sale"], "comparableLeases": ["demo-lease"],
}

DEMO_RETAIL_APPRAISAL = {
    "_id": "demo-market-hall", "owner": "local-demo", "name": "Market Hall Demo",
    "address": "45 Wellington Street, Toronto", "appraisalType": "detailed",
    "location": {"type": "Point", "coordinates": [-79.3732, 43.6481]},
    "propertyType": "retail", "sizeOfBuilding": 18_000, "effectiveDate": "2026-01-01",
    "units": [
        {"tenantName": "Harbour Bakery", "suite": "A", "unitNumber": "A", "squareFootage": 3_200,
         "yearlyRent": 128_000, "tenancies": [{"name": "Harbour Bakery", "yearlyRent": 128_000}]},
        {"tenantName": "Juniper Books", "suite": "B", "unitNumber": "B", "squareFootage": 2_100,
         "yearlyRent": 88_200, "tenancies": [{"name": "Juniper Books", "yearlyRent": 88_200}]},
    ],
    "additionalIncomes": [{"name": "Signage", "amount": 9_600}],
    "expenses": [{"name": "Common area maintenance", "amount": 42_000}, {"name": "Insurance", "amount": 8_400}],
    "marketRents": [{"name": "Retail market rent", "amountPSF": 42, "amount": 42}],
    "stabilizedStatementInputs": {"vacancyRate": 4, "capitalizationRate": 5.5, "structuralAllowancePercent": 1.5, "modifiers": []},
    "directComparisonInputs": {"directComparisonMetric": "psf", "pricePerSquareFoot": 425, "modifiers": []},
    "discountedCashFlowInputs": {"projectionYears": 10, "discountRate": 7.25, "inflation": 2},
    "comparableSalesCapRate": ["demo-sale-retail"], "comparableSalesDCA": ["demo-sale-retail"], "comparableLeases": ["demo-lease-retail"],
}


def with_appraisal_defaults(value: dict) -> dict:
    result = {**DEMO_DEFAULTS, **value}
    for key, default in DEMO_DEFAULTS.items():
        if isinstance(default, dict) and isinstance(value.get(key), dict):
            result[key] = {**default, **value[key]}
    return result


class JsonFormatter(logging.Formatter):
    """Small dependency-free formatter for logs consumed by Compose or CI."""

    def format(self, record: logging.LogRecord) -> str:
        payload = {
            "timestamp": datetime.now(UTC).isoformat(),
            "level": record.levelname,
            "logger": record.name,
            "message": record.getMessage(),
        }
        for key in ("appraisalId", "fileId", "jobId"):
            if value := getattr(record, key, None):
                payload[key] = value
        if record.exc_info:
            payload["exception"] = self.formatException(record.exc_info)
        return json.dumps(payload, default=str)


def configure_logging() -> None:
    handler = logging.StreamHandler()
    handler.setFormatter(JsonFormatter())
    log.handlers = [handler]
    log.setLevel(logging.INFO)
    log.propagate = False


def public(document: dict) -> dict:
    document = dict(document)
    document["_id"] = str(document.pop("_id"))
    document.pop("annotations", None)
    if isinstance(document.get("words"), list):
        document["words"] = [
            {key: value for key, value in word.items() if key in PUBLIC_WORD_FIELDS}
            if isinstance(word, dict) else word
            for word in document["words"]
        ]
    return document


def clean_payload(payload: dict) -> dict:
    """Legacy clients send Mongo `_id`; identity is always route-controlled."""
    return {key: value for key, value in payload.items() if key not in {"_id", "id", "owner", "createdAt"}}


def numeric_filters(params, mapping: dict[str, str]) -> dict:
    query: dict = {}
    for parameter, field in mapping.items():
        if value := params.get(f"{parameter}From"):
            query.setdefault(field, {})["$gte"] = float(value)
        if value := params.get(f"{parameter}To"):
            query.setdefault(field, {})["$lte"] = float(value)
    return query


def fixture_directory() -> Path:
    return Path(__file__).parents[1] / "fixtures"


def seed_demo_files(app: FastAPI) -> None:
    """Copy versioned sample documents into the persistent local file volume once."""
    cfg, db = settings(), app.state.db
    documents = (
        ("demo-financial-statement", "financial-statement.pdf", "financial-statement.json", "financials"),
        ("demo-lease", "lease.docx", "lease.json", "lease"),
        ("demo-comparable-sale", "comparable-sale.txt", "comparable-sale.json", "comparable"),
        ("demo-scanned-rent-roll", "scanned-placeholder.pdf", "partial.json", "rentroll"),
    )

    def word_records(values: list[object]) -> list[object]:
        """Keep persisted early-demo string tokens usable for source highlighting."""
        return [
            {"word": value, "page": 1, "index": index, "lineNumber": 0, "documentLineNumber": 0,
             "column": index, "documentColumn": index, "left": 0, "right": 0, "top": 0, "bottom": 0}
            if isinstance(value, str) else value
            for index, value in enumerate(values)
        ]

    for file_id, source_name, extraction_name, legacy_type in documents:
        existing = db.files.find_one({"_id": file_id})
        if existing:
            words = existing.get("words") or []
            if any(isinstance(word, str) for word in words):
                db.files.update_one({"_id": file_id}, {"$set": {"words": word_records(words)}})
            continue
        source = fixture_directory() / source_name
        if not source.is_file():
            log.warning("seed fixture is unavailable", extra={"fileId": file_id})
            continue
        destination = cfg.data_dir / "uploads" / f"{file_id}{source.suffix}"
        shutil.copy2(source, destination)
        extracted_text, pages = extract_available_text(destination, source.name)
        result = ExtractionResult.model_validate_json((fixture_directory() / extraction_name).read_text())
        normalized = normalize_extraction(result)
        normalized["fileType"] = legacy_type
        images = render_document_pages(file_id, destination)
        db.files.insert_one({
            "_id": file_id, "appraisalId": "demo-appraisal", "owner": "local-demo", "fileName": source_name,
            "path": str(destination), "hash": hashlib.sha256(destination.read_bytes()).hexdigest(),
            "reviewStatus": "seeded", "pages": pages, "images": [str(image) for image in images],
            "words": word_records(extracted_text.split()), **normalized,
        })


def seed_or_upgrade_demo(app: FastAPI) -> None:
    db = app.state.db
    demo = db.appraisals.find_one({"_id": "demo-appraisal"})
    if not demo:
        demo = with_appraisal_defaults({**DEMO_APPRAISAL, "createdAt": datetime.now(UTC), "updatedAt": datetime.now(UTC)})
        demo.update(refresh_valuations(demo))
        demo["demoSeedVersion"] = DEMO_SEED_VERSION
        db.appraisals.insert_one(demo)
    elif demo.get("demoSeedVersion") != DEMO_SEED_VERSION:
        # Upgrade the original minimal demo without replacing meaningful local edits.
        additions, merged = {}, dict(demo)
        if demo.get("appraisalType") not in {"simple", "detailed"}:
            additions["appraisalType"] = "detailed"
            merged["appraisalType"] = "detailed"
        if isinstance(demo.get("propertyType"), str) and demo["propertyType"].lower() in {"office", "industrial", "retail", "land", "residential"}:
            additions["propertyType"] = demo["propertyType"].lower()
            merged["propertyType"] = demo["propertyType"].lower()
        for key, value in DEMO_APPRAISAL.items():
            if key == "_id":
                continue
            current = demo.get(key)
            if key not in demo or current in (None, "", [], {}):
                additions[key] = value
                merged[key] = value
            elif isinstance(value, dict) and isinstance(current, dict):
                missing = {field: default for field, default in value.items()
                           if field not in current or current[field] in (None, "", [], {})}
                if missing:
                    merged[key] = {**current, **missing}
                    additions.update({f"{key}.{field}": default for field, default in missing.items()})
        # Add the current React screen-model fields without replacing local values.
        if demo.get("units"):
            upgraded_units = []
            default_units = {unit["unitNumber"]: unit for unit in DEMO_APPRAISAL["units"]}
            market_rents = merged.get("marketRents") or []
            default_market_rent = (market_rents[0] or {}).get("name") if market_rents else None
            for unit in demo["units"]:
                rent = unit.get("yearlyRent", unit.get("annualRent", 0))
                is_vacant = unit.get("isVacantForStabilizedStatement", unit.get("shouldTreatAsVacant", False))
                upgraded = {**unit, "unitNumber": unit.get("unitNumber", unit.get("suite", "")),
                            "shouldTreatAsVacant": is_vacant}
                default_unit = default_units.get(upgraded["unitNumber"], {})
                if not unit.get("tenancies"):
                    upgraded["tenancies"] = default_unit.get("tenancies") or [
                        {"name": unit.get("tenantName", "Vacant"), "yearlyRent": rent}
                    ]
                elif default_unit.get("tenancies"):
                    tenancy_defaults = default_unit["tenancies"]
                    upgraded["tenancies"] = [
                        {
                            **tenancy,
                            **{field: value for field, value in tenancy_defaults[min(index, len(tenancy_defaults) - 1)].items()
                               if field not in tenancy or tenancy[field] in (None, "", [], {})},
                        }
                        for index, tenancy in enumerate(unit["tenancies"])
                    ]
                if is_vacant and isinstance(unit.get("marketRent"), (int, float)) and default_market_rent:
                    upgraded.update({"marketRent": default_market_rent, "shouldUseMarketRent": True})
                upgraded_units.append(upgraded)
            merged["units"] = upgraded_units
            additions["units"] = upgraded_units
            merged["sizeOfBuilding"] = sum(unit.get("squareFootage", 0) for unit in upgraded_units)
            additions["sizeOfBuilding"] = merged["sizeOfBuilding"]
        if demo.get("marketRents"):
            upgraded_market_rents = [{**market_rent, "amountPSF": market_rent.get("amountPSF", market_rent.get("amount", 0))}
                                    for market_rent in demo["marketRents"]]
            merged["marketRents"] = upgraded_market_rents
            additions["marketRents"] = upgraded_market_rents
        for statement_name in ("incomeStatement", "expenseStatement"):
            for field_path in [key for key in additions if key.startswith(f"{statement_name}.")]:
                del additions[field_path]
            statement = dict(merged.get(statement_name) or {})
            defaults = DEMO_APPRAISAL[statement_name]
            items = list(statement.get("items") or [])
            for default_item in defaults["items"]:
                existing = next((item for item in items if item.get("name") == default_item["name"]), None)
                if existing is None:
                    items.append(default_item)
                else:
                    for field, value in default_item.items():
                        if field not in existing or existing[field] in (None, "", [], {}):
                            existing[field] = value
            statement.update({**defaults, **statement, "years": statement.get("years") or defaults["years"], "items": items})
            merged[statement_name] = statement
            additions[statement_name] = statement
        for field_path in [key for key in additions if key.startswith("amortizationSchedule.")]:
            del additions[field_path]
        schedule = dict(merged.get("amortizationSchedule") or {})
        schedule_items = list(schedule.get("items") or [])
        for default_item in DEMO_APPRAISAL["amortizationSchedule"]["items"]:
            existing = next((item for item in schedule_items if item.get("name") == default_item["name"]), None)
            if existing is None:
                schedule_items.append(default_item)
            else:
                for field, value in default_item.items():
                    if field not in existing or existing[field] in (None, "", [], {}):
                        existing[field] = value
        schedule["items"] = schedule_items
        merged["amortizationSchedule"] = schedule
        additions["amortizationSchedule"] = schedule
        additions.update(refresh_valuations(merged))
        additions["demoSeedVersion"] = DEMO_SEED_VERSION
        db.appraisals.update_one({"_id": "demo-appraisal"}, {"$set": additions})
    if not db.appraisals.find_one({"_id": DEMO_RETAIL_APPRAISAL["_id"]}):
        retail = with_appraisal_defaults({**DEMO_RETAIL_APPRAISAL, "createdAt": datetime.now(UTC), "updatedAt": datetime.now(UTC)})
        retail.update(refresh_valuations(retail))
        retail["demoSeedVersion"] = DEMO_SEED_VERSION
        db.appraisals.insert_one(retail)
    seed_demo_files(app)


def ensure_demo_comparables(db) -> None:
    """Add missing seeded comparables without replacing a local user's edits."""
    sales = (
        {"_id": "demo-sale", "owner": "local-demo", "address": "25 King Street, Toronto", "salePrice": 12_500_000,
         "sizeSquareFootage": 50_000, "saleDate": "2025-01-15", "propertyType": "Office", "allowSubCompSearch": True,
         "netOperatingIncome": 656_250, "capitalizationRate": 5.25, "occupancyRate": 94,
         "location": {"type": "Point", "coordinates": [-79.3791, 43.6487]}},
        {"_id": "demo-sale-retail", "owner": "local-demo", "address": "55 Queen Street, Toronto", "salePrice": 7_650_000,
         "sizeSquareFootage": 18_400, "saleDate": "2025-04-18", "propertyType": "Retail", "allowSubCompSearch": True,
         "netOperatingIncome": 428_400, "capitalizationRate": 5.6, "occupancyRate": 98,
         "location": {"type": "Point", "coordinates": [-79.3784, 43.6524]}},
    )
    leases = (
        {"_id": "demo-lease", "owner": "local-demo", "tenantName": "Northstar Foods", "address": "12 Front Street, Toronto",
         "sizeOfUnit": 8_000, "leaseDate": "2025-02-01", "propertyType": "Industrial",
         "rentEscalations": [{"startYear": 1, "endYear": 5, "yearlyRent": 160_000}],
         "taxesMaintenanceInsurance": 12.5, "tenantInducements": "$20 psf allowance", "freeRentMonths": 3,
         "freeRentType": "net", "location": {"type": "Point", "coordinates": [-79.3745, 43.6455]}},
        {"_id": "demo-lease-retail", "owner": "local-demo", "tenantName": "Harbour Bakery", "address": "51 Queen Street, Toronto",
         "sizeOfUnit": 3_100, "leaseDate": "2025-03-15", "propertyType": "Retail",
         "rentEscalations": [{"startYear": 1, "endYear": 5, "yearlyRent": 130_200}],
         "taxesMaintenanceInsurance": 18.75, "tenantInducements": "$35 psf allowance", "freeRentMonths": 2,
         "freeRentType": "net", "location": {"type": "Point", "coordinates": [-79.3778, 43.6522]}},
    )
    for record in sales:
        db.comparable_sales.update_one({"_id": record["_id"]}, {"$setOnInsert": record}, upsert=True)
        for field, value in record.items():
            if field not in {"_id", "owner"}:
                db.comparable_sales.update_one({"_id": record["_id"], field: {"$exists": False}}, {"$set": {field: value}})
    for record in leases:
        db.comparable_leases.update_one({"_id": record["_id"]}, {"$setOnInsert": record}, upsert=True)
        for field, value in record.items():
            if field not in {"_id", "owner", "rentEscalations"}:
                db.comparable_leases.update_one({"_id": record["_id"], field: {"$exists": False}}, {"$set": {field: value}})
        for field in ("startYear", "endYear"):
            db.comparable_leases.update_one(
                {"_id": record["_id"], f"rentEscalations.0.{field}": {"$exists": False}},
                {"$set": {f"rentEscalations.0.{field}": record["rentEscalations"][0][field]}},
            )


def ensure_demo_zone(db) -> None:
    """Keep the seeded appraisal's zoning reference resolvable in the editor."""
    db.zones.update_one(
        {"_id": "CR 3.0"},
        {"$set": {
            "owner": "local-demo",
            "zoneName": "CR 3.0",
            "description": "Commercial Residential zoning supporting a mixed-use office property.",
        }},
        upsert=True,
    )


@asynccontextmanager
async def lifespan(app: FastAPI):
    configure_logging()
    cfg = settings()
    app.state.mongo = MongoClient(cfg.mongo_url)
    app.state.db = app.state.mongo[cfg.mongo_db]
    app.state.db.files.create_index("appraisalId")
    app.state.db.extractions.create_index("fileId")
    seed_or_upgrade_demo(app)
    ensure_demo_comparables(app.state.db)
    ensure_demo_zone(app.state.db)
    log.info("api started", extra={"appraisalId": "local-demo"})
    yield
    app.state.mongo.close()
    log.info("api stopped")


app = FastAPI(title="Swiftly local demo API", version="1.0.0", lifespan=lifespan)
app.add_middleware(CORSMiddleware, allow_origin_regex=r"https?://(localhost|127\.0\.0\.1)(:\d+)?", allow_methods=["*"], allow_headers=["*"])


@app.get("/health")
def health(request: Request):
    request.app.state.mongo.admin.command("ping")
    return {"status": "ok"}


@app.get("/appraisal/")
@app.get("/appraisals")
def list_appraisals(request: Request):
    return {"appraisals": [public(item) for item in request.app.state.db.appraisals.find({}, {"name": 1, "address": 1, "appraisalType": 1})]}


@app.post("/appraisal/")
def create_appraisal(payload: dict, request: Request):
    appraisal_id = str(uuid4())
    payload = with_appraisal_defaults(clean_payload(payload))
    payload.update({"_id": appraisal_id, "owner": "local-demo", "createdAt": datetime.now(UTC), "updatedAt": datetime.now(UTC)})
    request.app.state.db.appraisals.insert_one(payload)
    return {"_id": appraisal_id}


@app.get("/appraisal/{appraisal_id}")
def get_appraisal(appraisal_id: str, request: Request):
    appraisal = request.app.state.db.appraisals.find_one({"_id": appraisal_id})
    if not appraisal: raise HTTPException(404, "Appraisal not found")
    return {"appraisal": public(appraisal)}


@app.post("/appraisal/{appraisal_id}")
def update_appraisal(appraisal_id: str, payload: dict, request: Request):
    payload = clean_payload(payload)
    existing = request.app.state.db.appraisals.find_one({"_id": appraisal_id})
    if not existing: raise HTTPException(404, "Appraisal not found")
    merged = {**existing, **payload}
    calculated = refresh_valuations(merged)
    result = request.app.state.db.appraisals.find_one_and_update({"_id": appraisal_id}, {"$set": {**payload, **calculated, "updatedAt": datetime.now(UTC)}}, return_document=ReturnDocument.AFTER)
    if not result: raise HTTPException(404, "Appraisal not found")
    return {"appraisal": public(result)}


@app.delete("/appraisal/{appraisal_id}")
def delete_appraisal(appraisal_id: str, request: Request):
    result = request.app.state.db.appraisals.find_one_and_delete({"_id": appraisal_id})
    if not result: raise HTTPException(404, "Appraisal not found")
    for record in request.app.state.db.files.find({"appraisalId": appraisal_id}):
        remove_file_assets(record)
    request.app.state.db.files.delete_many({"appraisalId": appraisal_id})
    request.app.state.db.extractions.delete_many({"appraisalId": appraisal_id})
    return {}


@app.post("/appraisal/{appraisal_id}/convert_tenants")
def convert_tenants(appraisal_id: str, request: Request):
    appraisal = request.app.state.db.appraisals.find_one({"_id": appraisal_id})
    if not appraisal: raise HTTPException(404, "Appraisal not found")
    created = []
    for unit in appraisal.get("units", []):
        if unit_is_vacant(unit): continue
        comparable = {**unit, "_id": str(uuid4()), "owner": "local-demo", "appraisalId": appraisal_id,
                      "tenantName": unit_tenant_name(unit), "propertyType": appraisal.get("propertyType")}
        request.app.state.db.comparable_leases.insert_one(comparable); created.append(comparable["_id"])
    return {"created": created}


def child_records(request: Request, collection: str, appraisal_id: str, response_key: str):
    return {response_key: [public(item) for item in request.app.state.db[collection].find({"appraisalId": appraisal_id})]}


def create_child(payload: dict, request: Request, collection: str, appraisal_id: str) -> dict:
    record = {**clean_payload(payload), "_id": str(uuid4()), "appraisalId": appraisal_id, "owner": "local-demo"}
    request.app.state.db[collection].insert_one(record)
    return {"_id": record["_id"]}


def get_child(request: Request, collection: str, appraisal_id: str, record_id: str, response_key: str):
    record = request.app.state.db[collection].find_one({"_id": record_id, "appraisalId": appraisal_id})
    if not record: raise HTTPException(404, "Record not found")
    return {response_key: public(record)}


def update_child(payload: dict, request: Request, collection: str, appraisal_id: str, record_id: str) -> dict:
    result = request.app.state.db[collection].update_one({"_id": record_id, "appraisalId": appraisal_id}, {"$set": clean_payload(payload)})
    if not result.matched_count: raise HTTPException(404, "Record not found")
    return {"_id": record_id}


@app.get("/appraisal/{appraisal_id}/leases")
def list_leases(appraisal_id: str, request: Request):
    return child_records(request, "leases", appraisal_id, "leases")


@app.post("/appraisal/{appraisal_id}/leases")
def create_lease(appraisal_id: str, payload: dict, request: Request):
    return create_child(payload, request, "leases", appraisal_id)


@app.get("/appraisal/{appraisal_id}/leases/{lease_id}")
def get_lease(appraisal_id: str, lease_id: str, request: Request):
    return get_child(request, "leases", appraisal_id, lease_id, "lease")


@app.post("/appraisal/{appraisal_id}/leases/{lease_id}")
def update_lease(appraisal_id: str, lease_id: str, payload: dict, request: Request):
    return update_child(payload, request, "leases", appraisal_id, lease_id)


@app.get("/appraisal/{appraisal_id}/financial_statements")
def list_financial_statements(appraisal_id: str, request: Request):
    return child_records(request, "financial_statements", appraisal_id, "financial_statements")


@app.post("/appraisal/{appraisal_id}/financial_statements")
def create_financial_statement(appraisal_id: str, payload: dict, request: Request):
    return create_child(payload, request, "financial_statements", appraisal_id)


@app.get("/appraisal/{appraisal_id}/financial_statements/{statement_id}")
def get_financial_statement(appraisal_id: str, statement_id: str, request: Request):
    return get_child(request, "financial_statements", appraisal_id, statement_id, "financialStatement")


@app.post("/appraisal/{appraisal_id}/financial_statements/{statement_id}")
def update_financial_statement(appraisal_id: str, statement_id: str, payload: dict, request: Request):
    return update_child(payload, request, "financial_statements", appraisal_id, statement_id)


def collection_list(request: Request, collection: str, response_key: str, query: dict | None = None, sort: str | None = None):
    records = request.app.state.db[collection].find(query or {})
    if sort: records = records.sort(sort.lstrip("-"), -1 if sort.startswith("-") else 1)
    return {response_key: [public(record) for record in records]}


@app.get("/comparable_sales")
def list_comparable_sales(request: Request):
    params = request.query_params
    query = numeric_filters(params, {"salePrice": "salePrice", "leaseableArea": "sizeSquareFootage", "capitalizationRate": "capitalizationRate", "pricePerSquareFoot": "pricePerSquareFoot", "clearCeilingHeight": "clearCeilingHeight", "shippingDoors": "shippingDoors", "siteCoverage": "siteCoverage", "sizeOfLandAcres": "sizeOfLandAcres", "sizeOfLandSqft": "sizeOfLandSqft", "pricePerSquareFootLand": "pricePerSquareFootLand", "pricePerAcreLand": "pricePerAcreLand", "pricePerSquareFootBuildableArea": "pricePerSquareFootBuildableArea"})
    if property_type := params.get("propertyType"): query["propertyType"] = property_type
    if tags := params.getlist("propertyTags[]"): query["propertyTags"] = {"$all": tags}
    return collection_list(request, "comparable_sales", "comparableSales", query, params.get("sort") or "-saleDate")


@app.post("/comparable_sales")
def create_comparable_sale(payload: dict, request: Request):
    record = {**clean_payload(payload), "_id": str(uuid4()), "owner": "local-demo", "createdAt": datetime.now(UTC)}
    request.app.state.db.comparable_sales.insert_one(record)
    return {"_id": record["_id"]}


@app.get("/comparable_sales/{comparable_id}")
def get_comparable_sale(comparable_id: str, request: Request):
    record = request.app.state.db.comparable_sales.find_one({"_id": comparable_id})
    return {"comparableSale": public(record) if record else None}


@app.post("/comparable_sales/{comparable_id}")
def update_comparable_sale(comparable_id: str, payload: dict, request: Request):
    result = request.app.state.db.comparable_sales.update_one({"_id": comparable_id}, {"$set": clean_payload(payload)})
    if not result.matched_count: raise HTTPException(404, "Comparable sale not found")
    return {"_id": comparable_id}


@app.delete("/comparable_sales/{comparable_id}")
@app.delete("/comparable_sale/{comparable_id}")
def delete_comparable_sale(comparable_id: str, request: Request):
    if not request.app.state.db.comparable_sales.delete_one({"_id": comparable_id}).deleted_count: raise HTTPException(404, "Comparable sale not found")
    return {}


@app.post("/comparable_sales_portfolio/")
def save_comparable_sale_portfolio(payload: dict, request: Request):
    """Persist the legacy portfolio edit dialog atomically as individual comps."""
    raw_portfolio = payload.get("portfolio") or {}
    portfolio = clean_payload(raw_portfolio)
    sub_comps = payload.get("subComps") or []
    portfolio_id = raw_portfolio.get("_id") or str(uuid4())
    portfolio.update({"_id": portfolio_id, "owner": "local-demo", "isPortfolio": True, "subCompIds": []})
    for raw_comp in sub_comps:
        comp = clean_payload(raw_comp); comp_id = raw_comp.get("_id") or str(uuid4())
        comp.update({"_id": comp_id, "owner": "local-demo", "portfolioId": portfolio_id})
        request.app.state.db.comparable_sales.replace_one({"_id": comp_id}, comp, upsert=True)
        portfolio["subCompIds"].append(comp_id)
    request.app.state.db.comparable_sales.replace_one({"_id": portfolio_id}, portfolio, upsert=True)
    return {"_id": portfolio_id, "subCompIds": portfolio["subCompIds"]}


@app.get("/comparable_leases")
def list_comparable_leases(request: Request):
    params = request.query_params
    query = numeric_filters(params, {"sizeOfUnit": "sizeOfUnit", "taxesMaintenanceInsurance": "taxesMaintenanceInsurance"})
    if property_type := params.get("propertyType"): query["propertyType"] = property_type
    if tenancy_type := params.get("tenancyType"): query["tenancyType"] = tenancy_type
    if tenant_name := params.get("tenantName"): query["tenantName"] = {"$regex": tenant_name, "$options": "i"}
    if tags := params.getlist("propertyTags[]"): query["propertyTags"] = {"$all": tags}
    return collection_list(request, "comparable_leases", "comparableLeases", query, params.get("sort") or "-leaseDate")


@app.post("/comparable_leases")
def create_comparable_lease(payload: dict, request: Request):
    record = {**clean_payload(payload), "_id": str(uuid4()), "owner": "local-demo", "createdAt": datetime.now(UTC)}
    request.app.state.db.comparable_leases.insert_one(record)
    return {"_id": record["_id"]}


@app.get("/comparable_leases/{comparable_id}")
def get_comparable_lease(comparable_id: str, request: Request):
    record = request.app.state.db.comparable_leases.find_one({"_id": comparable_id})
    if not record: raise HTTPException(404, "Comparable lease not found")
    return {"comparableLease": public(record)}


@app.post("/comparable_leases/{comparable_id}")
def update_comparable_lease(comparable_id: str, payload: dict, request: Request):
    result = request.app.state.db.comparable_leases.update_one({"_id": comparable_id}, {"$set": clean_payload(payload)})
    if not result.matched_count: raise HTTPException(404, "Comparable lease not found")
    return {"_id": comparable_id}


@app.delete("/comparable_leases/{comparable_id}")
def delete_comparable_lease(comparable_id: str, request: Request):
    if not request.app.state.db.comparable_leases.delete_one({"_id": comparable_id}).deleted_count: raise HTTPException(404, "Comparable lease not found")
    return {}


@app.get("/zones")
def list_zones(request: Request):
    query = {}
    if name := request.query_params.get("zoneName"): query["zoneName"] = {"$regex": name, "$options": "i"}
    return collection_list(request, "zones", "zones", query)


@app.post("/zones")
def create_zone(payload: dict, request: Request):
    record = {**clean_payload(payload), "_id": str(uuid4()), "owner": "local-demo"}
    request.app.state.db.zones.insert_one(record)
    return {"_id": record["_id"]}


@app.get("/zone/{zone_id}")
def get_zone(zone_id: str, request: Request):
    record = request.app.state.db.zones.find_one({"_id": zone_id})
    if not record: raise HTTPException(404, "Zone not found")
    return {"zone": public(record)}


@app.post("/zone/{zone_id}")
def update_zone(zone_id: str, payload: dict, request: Request):
    if not request.app.state.db.zones.update_one({"_id": zone_id}, {"$set": clean_payload(payload)}).matched_count: raise HTTPException(404, "Zone not found")
    return {"_id": zone_id}


@app.delete("/zone/{zone_id}")
def delete_zone(zone_id: str, request: Request):
    if not request.app.state.db.zones.delete_one({"_id": zone_id}).deleted_count: raise HTTPException(404, "Zone not found")
    return {}


@app.get("/property_tags")
def list_property_tags(request: Request):
    query = {}
    if name := request.query_params.get("name"): query["name"] = {"$regex": name, "$options": "i"}
    if property_type := request.query_params.get("propertyType"): query["propertyType"] = property_type
    records = request.app.state.db.property_tags.find(query).limit(10)
    return {"tags": [public(record) for record in records]}


@app.post("/property_tags")
def create_property_tag(payload: dict, request: Request):
    record = {**clean_payload(payload), "_id": str(uuid4()), "owner": "local-demo"}
    request.app.state.db.property_tags.insert_one(record)
    return {"_id": record["_id"]}


@app.delete("/property_tags/{tag_id}")
def delete_property_tag(tag_id: str, request: Request):
    deleted = request.app.state.db.property_tags.delete_one({"$or": [{"_id": tag_id}, {"name": tag_id}]}).deleted_count
    if not deleted: raise HTTPException(404, "Property tag not found")
    return {}


@app.get("/tenant_names")
def tenant_names(request: Request):
    query = {"tenantName": {"$ne": ""}}
    if name := request.query_params.get("tenantName"): query["tenantName"] = {"$regex": name, "$options": "i"}
    names = request.app.state.db.comparable_leases.distinct("tenantName", query)
    return {"names": sorted(name for name in names if name)}


IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".gif"}
COMPARABLE_IMPORT_FIELDS = {
    "address": "address", "saleprice": "salePrice", "sale_price": "salePrice",
    "salesprice": "salePrice", "sale_date": "saleDate", "saledate": "saleDate",
    "size_square_footage": "sizeSquareFootage", "sizesquarefootage": "sizeSquareFootage",
    "leaseable_area": "sizeSquareFootage", "leaseablearea": "sizeSquareFootage",
    "price_per_square_foot": "pricePerSquareFoot", "pricepersquarefoot": "pricePerSquareFoot",
    "capitalization_rate": "capitalizationRate", "capitalizationrate": "capitalizationRate",
    "property_type": "propertyType", "propertytype": "propertyType",
}
COMPARABLE_NUMERIC_FIELDS = {"salePrice", "sizeSquareFootage", "pricePerSquareFoot", "capitalizationRate"}


def image_url(request: Request, image_id: str) -> str:
    return f"{str(request.base_url).rstrip('/')}/images/{image_id}"


def save_image(upload: UploadFile, request: Request) -> dict:
    suffix = Path(upload.filename or "image").suffix.lower()
    if suffix not in IMAGE_SUFFIXES:
        raise HTTPException(415, "Images must be PNG, JPEG, WebP, or GIF files")
    image_id = str(uuid4())
    destination = settings().data_dir / "images" / f"{image_id}{suffix}"
    with destination.open("wb") as target:
        shutil.copyfileobj(upload.file, target)
    record = {"_id": image_id, "fileName": upload.filename or f"image{suffix}", "path": str(destination),
              "contentType": upload.content_type or "application/octet-stream", "createdAt": datetime.now(UTC)}
    request.app.state.db.images.insert_one(record)
    return record


@app.post("/images")
def upload_image(request: Request, file: UploadFile = Upload(...)):
    record = save_image(file, request)
    return {"url": image_url(request, record["_id"])}


@app.get("/images/{image_id}")
def get_image(image_id: str, request: Request):
    record = request.app.state.db.images.find_one({"_id": image_id})
    if not record or not Path(record["path"]).is_file():
        raise HTTPException(404, "Image not found")
    return FileResponse(record["path"], media_type=record.get("contentType"))


def imported_comparable_rows(upload: UploadFile) -> list[dict]:
    suffix = Path(upload.filename or "").suffix.lower()
    if suffix == ".json":
        try:
            decoded = json.load(upload.file)
        except json.JSONDecodeError as exc:
            raise HTTPException(422, "Comparable import JSON is invalid") from exc
        if isinstance(decoded, dict) and decoded.get("document_type") == "comparable_sale":
            fields = {item.get("name"): item.get("value") for item in decoded.get("fields", []) if isinstance(item, dict)}
            extracted_rows = decoded.get("comparable_sales") or []
            decoded = [
                {**fields, **{item.get("name"): item.get("value") for item in extracted.get("values", []) if isinstance(item, dict)}}
                for extracted in extracted_rows if isinstance(extracted, dict)
            ] or [fields]
        elif isinstance(decoded, dict):
            decoded = decoded.get("comparableSales", decoded.get("comparables", []))
        if not isinstance(decoded, list):
            raise HTTPException(422, "Comparable import JSON must contain a list of comparables")
        raw_rows = decoded
    elif suffix in {".csv", ".tsv"}:
        delimiter = "\t" if suffix == ".tsv" else ","
        raw_rows = list(csv.DictReader((line.decode("utf-8-sig") for line in upload.file), delimiter=delimiter))
    else:
        raise HTTPException(415, "Comparable imports must be CSV, TSV, or JSON files")

    rows: list[dict] = []
    for raw in raw_rows:
        if not isinstance(raw, dict):
            continue
        row = {"imageUrls": [], "captions": [], "propertyTags": []}
        for key, value in raw.items():
            field = COMPARABLE_IMPORT_FIELDS.get(str(key).strip().lower().replace(" ", "_"))
            if not field or value in (None, ""):
                continue
            if field in COMPARABLE_NUMERIC_FIELDS:
                try:
                    value = float(str(value).replace(",", "").replace("$", "").replace("%", ""))
                except ValueError:
                    continue
                value = int(value) if value.is_integer() else value
            row[field] = value
        if row.get("address"):
            rows.append(row)
    return rows


@app.post("/comparable_sale_upload/")
def upload_comparable_sales(request: Request, file: UploadFile = Upload(...)):
    rows = imported_comparable_rows(file)
    import_id = str(uuid4())
    request.app.state.db.imports.insert_one({"_id": import_id, "owner": "local-demo", "fileName": file.filename or "import",
                                              "kind": "comparable_sales", "rows": rows, "createdAt": datetime.now(UTC)})
    return {"file": {"_id": import_id, "fileName": file.filename or "import", "fileType": "comparable_sales_import", "pages": 1},
            "comparableSales": rows}


REPORT_TITLES = {
    "comparable_sales": "Comparable Sales", "comparable_leases": "Comparable Leases", "rent_roll": "Rent Roll",
    "expenses": "Expenses", "stabilized_statement": "Stabilized Statement", "capitalization_valuation": "Capitalization Valuation",
    "direct_comparison_valuation": "Direct Comparison Valuation", "additional_incomes": "Additional Income",
    "amortization_schedule": "Amortization Schedule", "tenants": "Tenants", "market_rents": "Market Rents", "subject_details": "Subject Details",
}


def linked_records(db, collection: str, values: list[object]) -> list[dict]:
    """Resolve legacy embedded records or modern ID lists without changing report order."""
    embedded = [value for value in values if isinstance(value, dict)]
    ids = [value for value in values if isinstance(value, str)]
    if not ids or db is None:
        return embedded
    records = {record["_id"]: record for record in db[collection].find({"_id": {"$in": ids}})}
    return embedded + [records[record_id] for record_id in ids if record_id in records]


def report_rows(appraisal: dict, report: str, db=None) -> list[tuple[str, object]]:
    if report == "comparable_sales":
        values = appraisal.get("comparableSales") or list(dict.fromkeys((appraisal.get("comparableSalesCapRate") or []) + (appraisal.get("comparableSalesDCA") or [])))
        return [(item.get("address", "Comparable"), item.get("salePrice", "")) for item in linked_records(db, "comparable_sales", values)]
    if report == "comparable_leases":
        return [(item.get("tenantName", "Lease"), item.get("sizeOfUnit", "")) for item in linked_records(db, "comparable_leases", appraisal.get("comparableLeases", []))]
    if report in {"rent_roll", "tenants"}: return [(unit_tenant_name(item), unit_yearly_rent(appraisal, item)) for item in appraisal.get("units", [])]
    if report == "expenses": return [(item.get("name", "Expense"), item.get("amount", item.get("yearlyAmount", ""))) for item in appraisal.get("expenses", [])]
    if report == "additional_incomes": return [(item.get("name", "Income"), item.get("amount", item.get("yearlyAmount", ""))) for item in appraisal.get("additionalIncomes", [])]
    if report == "amortization_schedule": return [(item.get("name", "Amortization"), item.get("amount", "")) for item in (appraisal.get("amortizationSchedule") or {}).get("items", [])]
    if report in {"stabilized_statement", "capitalization_valuation"}: return list((appraisal.get("stabilizedStatement") or {}).items())
    if report == "direct_comparison_valuation": return list((appraisal.get("directComparisonValuation") or {}).items())
    return [("Property", appraisal.get("name", "")), ("Address", appraisal.get("address", ""))]


@app.get("/appraisal/{appraisal_id}/files")
def list_files(appraisal_id: str, request: Request, type: str | None = None):
    query = {"appraisalId": appraisal_id}
    if type: query["fileType"] = type
    fields = {"fileName": 1, "fileType": 1, "reviewStatus": 1, "extractionJobId": 1, "extractionError": 1, "pages": 1}
    return {"files": [public(item) for item in request.app.state.db.files.find(query, fields)]}


def save_upload(appraisal_id: str, upload: UploadFile, request: Request) -> dict:
    file_id = str(uuid4()); cfg = settings(); suffix = Path(upload.filename or "upload").suffix
    destination = cfg.data_dir / "uploads" / f"{file_id}{suffix}"
    with destination.open("wb") as target: shutil.copyfileobj(upload.file, target)
    digest = hashlib.sha256(destination.read_bytes()).hexdigest()
    record = {"_id": file_id, "appraisalId": appraisal_id, "owner": "local-demo", "fileName": upload.filename or "upload",
              "fileType": "other", "reviewStatus": "fresh", "path": str(destination), "hash": digest, "images": [], "words": [], "pages": 0}
    request.app.state.db.files.insert_one(record)
    return record


def render_document_pages(file_id: str, source: Path) -> list[Path]:
    """Persist reviewable PNGs; OCR-free scanned documents are still sent visually."""
    cfg = settings(); output_dir = cfg.data_dir / "rendered" / file_id; output_dir.mkdir(parents=True, exist_ok=True)
    if source.suffix.lower() != ".pdf":
        return [source] if source.suffix.lower() in {".png", ".jpg", ".jpeg", ".webp"} else []
    document = fitz.open(source); images = []
    for index, page in enumerate(document):
        destination = output_dir / f"page-{index + 1}.png"
        page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False).save(destination)
        images.append(destination)
    return images


def extract_available_text(source: Path, file_name: str) -> tuple[str, int]:
    """Extract locally available text before sending source pages to the provider."""
    suffix = source.suffix.lower()
    if suffix == ".pdf":
        document = fitz.open(source)
        return "\n\n".join(f"[Page {i + 1}]\n{page.get_text()}" for i, page in enumerate(document)), document.page_count
    if suffix == ".docx":
        document = Document(source)
        return "[Page 1]\n" + "\n".join(paragraph.text for paragraph in document.paragraphs), 1
    if suffix in {".csv", ".txt", ".tsv"}:
        return "[Page 1]\n" + source.read_text(encoding="utf-8", errors="replace"), 1
    return f"[Page 1] Image upload: {file_name}", 1


def queue_extraction(appraisal_id: str, file_id: str, request: Request, background_tasks: BackgroundTasks) -> str:
    """Create a reviewable job and keep the file's visible extraction state in sync."""
    job_id = str(uuid4())
    request.app.state.db.extractions.insert_one({"_id": job_id, "fileId": file_id, "appraisalId": appraisal_id,
                                                  "status": "queued", "createdAt": datetime.now(UTC)})
    request.app.state.db.files.update_one({"_id": file_id, "appraisalId": appraisal_id}, {"$set": {
        "extractionJobId": job_id, "reviewStatus": "queued", "extractionError": None,
    }})
    background_tasks.add_task(run_extraction, job_id, request)
    return job_id


def remove_file_assets(record: dict) -> None:
    """Remove only files owned by this record from the local persistent volume."""
    cfg = settings()
    upload_dir = (cfg.data_dir / "uploads").resolve()
    source = Path(record["path"]).resolve()
    if source.is_relative_to(upload_dir):
        source.unlink(missing_ok=True)
    rendered_dir = cfg.data_dir / "rendered" / str(record["_id"])
    shutil.rmtree(rendered_dir, ignore_errors=True)


@app.post("/appraisals/{appraisal_id}/files")
def upload_file(appraisal_id: str, request: Request, background_tasks: BackgroundTasks, file: UploadFile = Upload(...)):
    if not request.app.state.db.appraisals.find_one({"_id": appraisal_id}):
        raise HTTPException(404, "Appraisal not found")
    record = save_upload(appraisal_id, file, request)
    if Path(record["fileName"]).suffix.lower() in {".pdf", ".png", ".jpg", ".jpeg", ".webp"}:
        job_id = queue_extraction(appraisal_id, record["_id"], request, background_tasks)
        record["extractionJobId"] = job_id
        record["reviewStatus"] = "queued"
    return {"file": public(record)}


@app.post("/appraisal/{appraisal_id}/files")
def legacy_upload_file(appraisal_id: str, request: Request, background_tasks: BackgroundTasks, file: UploadFile = Upload(...)):
    return upload_file(appraisal_id, request, background_tasks, file)


def run_extraction(job_id: str, request: Request):
    db = request.app.state.db; job = db.extractions.find_one({"_id": job_id}); file = db.files.find_one({"_id": job["fileId"]})
    db.extractions.update_one({"_id": job_id}, {"$set": {"status": "running", "startedAt": datetime.now(UTC)}})
    try:
        path = Path(file["path"]); page_images = render_document_pages(file["_id"], path)
        text, file["pages"] = extract_available_text(path, file["fileName"])
        # Rendering and local text extraction are valuable even when the provider
        # is unavailable. Persist them before the remote call so a reviewer can
        # inspect the source and correct fields manually after a failed job.
        db.files.update_one({"_id": file["_id"]}, {"$set": {
            "pages": file["pages"], "images": [str(image) for image in page_images],
        }})
        result = provider_for(settings()).extract(path, text, page_images)
        normalized = normalize_extraction(result)
        db.files.update_one({"_id": file["_id"]}, {
            "$set": {**normalized, "reviewStatus": "extracted", "extractionError": None,
                     "pages": file["pages"], "images": [str(image) for image in page_images]},
            "$unset": {"annotations": ""},
        })
        db.extractions.update_one({"_id": job_id}, {"$set": {"status": "completed", "result": result.model_dump(), "completedAt": datetime.now(UTC)}})
    except Exception as exc:
        log.exception("extraction failed", extra={"jobId": job_id, "fileId": file["_id"], "appraisalId": job["appraisalId"]})
        db.extractions.update_one({"_id": job_id}, {"$set": {"status": "failed", "error": str(exc), "completedAt": datetime.now(UTC)}})
        db.files.update_one({"_id": file["_id"]}, {"$set": {"reviewStatus": "extraction_failed", "extractionError": str(exc)}})


@app.post("/appraisals/{appraisal_id}/files/{file_id}/extract")
@app.post("/appraisal/{appraisal_id}/files/{file_id}/reprocess")
def extract_file(appraisal_id: str, file_id: str, request: Request, background_tasks: BackgroundTasks):
    if not request.app.state.db.files.find_one({"_id": file_id, "appraisalId": appraisal_id}): raise HTTPException(404, "File not found")
    job_id = queue_extraction(appraisal_id, file_id, request, background_tasks)
    return {"jobId": job_id, "status": "queued"}


@app.get("/extractions/{job_id}")
def get_extraction(job_id: str, request: Request):
    job = request.app.state.db.extractions.find_one({"_id": job_id})
    if not job: raise HTTPException(404, "Extraction job not found")
    return {"extraction": public(job)}


@app.patch("/appraisals/{appraisal_id}/files/{file_id}/extraction")
def patch_extraction(appraisal_id: str, file_id: str, payload: ExtractionPatch, request: Request):
    normalized = normalize_extraction(payload.extraction)
    result = request.app.state.db.files.find_one_and_update(
        {"_id": file_id, "appraisalId": appraisal_id},
        {"$set": {**normalized, "reviewStatus": "corrected"}, "$unset": {"annotations": ""}},
        return_document=ReturnDocument.AFTER,
    )
    if not result: raise HTTPException(404, "File not found")
    return {"file": public(result)}


@app.get("/appraisal/{appraisal_id}/files/{file_id}")
def get_file(appraisal_id: str, file_id: str, request: Request):
    record = request.app.state.db.files.find_one({"_id": file_id, "appraisalId": appraisal_id})
    if not record: raise HTTPException(404, "File not found")
    return {"file": public(record)}


@app.post("/appraisal/{appraisal_id}/files/{file_id}")
def update_file(appraisal_id: str, file_id: str, payload: dict, request: Request):
    payload.pop("_id", None)
    record = request.app.state.db.files.find_one_and_update({"_id": file_id, "appraisalId": appraisal_id}, {"$set": payload}, return_document=ReturnDocument.AFTER)
    if not record: raise HTTPException(404, "File not found")
    return {"file": public(record)}


@app.delete("/appraisal/{appraisal_id}/files/{file_id}")
def delete_file(appraisal_id: str, file_id: str, request: Request):
    record = request.app.state.db.files.find_one_and_delete({"_id": file_id, "appraisalId": appraisal_id})
    if not record: raise HTTPException(404, "File not found")
    remove_file_assets(record)
    request.app.state.db.extractions.delete_many({"fileId": file_id})
    return {}


@app.get("/appraisal/{appraisal_id}/files/{file_id}/contents")
def file_contents(appraisal_id: str, file_id: str, request: Request):
    record = request.app.state.db.files.find_one({"_id": file_id, "appraisalId": appraisal_id})
    if not record: raise HTTPException(404, "File not found")
    return FileResponse(record["path"], filename=record["fileName"])


@app.get("/appraisal/{appraisal_id}/files/{file_id}/rendered/{page}")
def rendered_file_page(appraisal_id: str, file_id: str, page: int, request: Request):
    record = request.app.state.db.files.find_one({"_id": file_id, "appraisalId": appraisal_id})
    if not record: raise HTTPException(404, "File not found")
    images = record.get("images", [])
    if page < 1 or page > len(images): raise HTTPException(404, "Rendered page not found")
    return FileResponse(images[page - 1])


# This broad legacy route is intentionally declared after specific file routes.
# Starlette matches in declaration order; putting it earlier would swallow file GETs.
@app.get("/appraisal/{appraisal_id}/{report}/{format}")
def export_report(appraisal_id: str, report: str, format: str, request: Request):
    if report not in REPORT_TITLES or format not in {"word", "excel", "detailed_word"}: raise HTTPException(404, "Export not found")
    appraisal = request.app.state.db.appraisals.find_one({"_id": appraisal_id})
    if not appraisal: raise HTTPException(404, "Appraisal not found")
    title, rows = REPORT_TITLES[report], report_rows(appraisal, report, request.app.state.db)
    if format in {"word", "detailed_word"}:
        document = Document(); document.add_heading(title, 0); document.add_paragraph(appraisal.get("name", ""))
        table = document.add_table(rows=1, cols=2); table.style = "Light Shading Accent 1"; table.rows[0].cells[0].text = "Item"; table.rows[0].cells[1].text = "Value"
        for label, value in rows: cells = table.add_row().cells; cells[0].text = str(label); cells[1].text = str(value)
        output = BytesIO(); document.save(output); output.seek(0)
        return StreamingResponse(output, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f'attachment; filename="{report}.docx"'})
    workbook = Workbook(); sheet = workbook.active; sheet.title = title[:31]; sheet.append(["Item", "Value"])
    for row in rows: sheet.append(list(row))
    output = BytesIO(); workbook.save(output); output.seek(0)
    return StreamingResponse(output, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": f'attachment; filename="{report}.xlsx"'})
