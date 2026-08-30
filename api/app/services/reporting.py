"""Report row mapping and document generation independent from HTTP routing."""
from __future__ import annotations

from io import BytesIO

from docx import Document
from fastapi import HTTPException
from fastapi.responses import StreamingResponse
from openpyxl import Workbook

from ..calculations import unit_tenant_name, unit_yearly_rent

REPORT_TITLES = {
    "comparable_sales": "Comparable Sales",
    "comparable_leases": "Comparable Leases",
    "rent_roll": "Rent Roll",
    "expenses": "Expenses",
    "stabilized_statement": "Stabilized Statement",
    "capitalization_valuation": "Capitalization Valuation",
    "direct_comparison_valuation": "Direct Comparison Valuation",
    "additional_incomes": "Additional Income",
    "amortization_schedule": "Amortization Schedule",
    "tenants": "Tenants",
    "market_rents": "Market Rents",
    "subject_details": "Subject Details",
}


def linked_records(db, collection: str, values: list[object]) -> list[dict]:
    """Resolve embedded legacy records or ID lists without changing report order."""
    embedded = [value for value in values if isinstance(value, dict)]
    ids = [value for value in values if isinstance(value, str)]
    if not ids or db is None:
        return embedded
    records = {record["_id"]: record for record in db[collection].find({"_id": {"$in": ids}})}
    return embedded + [records[record_id] for record_id in ids if record_id in records]


def report_rows(appraisal: dict, report: str, db=None) -> list[tuple[str, object]]:
    if report == "comparable_sales":
        values = appraisal.get("comparableSales") or list(
            dict.fromkeys(
                (appraisal.get("comparableSalesCapRate") or [])
                + (appraisal.get("comparableSalesDCA") or [])
            )
        )
        return [
            (item.get("address", "Comparable"), item.get("salePrice", ""))
            for item in linked_records(db, "comparable_sales", values)
        ]
    if report == "comparable_leases":
        return [
            (item.get("tenantName", "Lease"), item.get("sizeOfUnit", ""))
            for item in linked_records(db, "comparable_leases", appraisal.get("comparableLeases", []))
        ]
    if report in {"rent_roll", "tenants"}:
        return [
            (unit_tenant_name(item), unit_yearly_rent(appraisal, item))
            for item in appraisal.get("units", [])
        ]
    if report == "expenses":
        return [
            (item.get("name", "Expense"), item.get("amount", item.get("yearlyAmount", "")))
            for item in appraisal.get("expenses", [])
        ]
    if report == "additional_incomes":
        return [
            (item.get("name", "Income"), item.get("amount", item.get("yearlyAmount", "")))
            for item in appraisal.get("additionalIncomes", [])
        ]
    if report == "amortization_schedule":
        return [
            (item.get("name", "Amortization"), item.get("amount", ""))
            for item in (appraisal.get("amortizationSchedule") or {}).get("items", [])
        ]
    if report in {"stabilized_statement", "capitalization_valuation"}:
        return list((appraisal.get("stabilizedStatement") or {}).items())
    if report == "direct_comparison_valuation":
        return list((appraisal.get("directComparisonValuation") or {}).items())
    return [("Property", appraisal.get("name", "")), ("Address", appraisal.get("address", ""))]


def export_report_response(appraisal: dict, report: str, output_format: str, db):
    format_map = {"docx": "word", "xlsx": "excel", "detailed-docx": "detailed_word"}
    legacy_format = format_map.get(output_format)
    if report not in REPORT_TITLES or not legacy_format:
        raise HTTPException(404, "Export not found")
    title = REPORT_TITLES[report]
    rows = report_rows(appraisal, report, db)
    if legacy_format in {"word", "detailed_word"}:
        document = Document()
        document.add_heading(title, 0)
        document.add_paragraph(appraisal.get("name", ""))
        table = document.add_table(rows=1, cols=2)
        table.style = "Light Shading Accent 1"
        table.rows[0].cells[0].text = "Item"
        table.rows[0].cells[1].text = "Value"
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = str(label)
            cells[1].text = str(value)
        output = BytesIO()
        document.save(output)
        output.seek(0)
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{report}.docx"'},
        )
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = title[:31]
    sheet.append(["Item", "Value"])
    for row in rows:
        sheet.append(list(row))
    output = BytesIO()
    workbook.save(output)
    output.seek(0)
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{report}.xlsx"'},
    )
